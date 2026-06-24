"""
memo_builder.py  --  builds the formal BOC "Daily Collection Report and
Assessment Import Analysis" memorandum (sections A-F) from a computed context,
in two renderings that share ONE block list so the narrative is identical:
  * narrative_html(ctx)  -> HTML for the dashboard's Memo tab
  * build_docx(ctx)      -> .docx bytes in the same layout/format as the sample

The phrasing mirrors the sample memo; numbers are computed live by app.py.
"""

import io
import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Header officials (from the sample memo; edit here if they change) ────────
OFFICIALS = {
    "ref": "ASMT-06-____",
    "for_name": "CARMELITA M. TALUSAN, CESO V",
    "for_title": "District Collector, Port of Batangas",
    "thru_name": "JULIUS HERALD B. ALTICHE",
    "thru_title": "Deputy Collector for Assessment",
    "from_name": "ATTY. ADRIAN M. VALENZUELA",
    "from_title": "Acting Chief, Assessment Division",
    "subject": "Daily Collection Report and Assessment Import Analysis",
}

NAVY = "16365C"
LIGHT = "DCE6F1"


# ── Formatting ───────────────────────────────────────────────────────────────

def php(x):
    if x is None:
        return "n.a."
    a = abs(x)
    if a >= 1e9:
        return f"Php{x/1e9:,.3f} billion"
    if a >= 1e6:
        return f"Php{x/1e6:,.3f} million"
    return f"Php{x:,.2f}"


def money(x):
    return "—" if x is None else f"{x:,.2f}"


def pct(x, plus=True):
    if x is None:
        return "new"
    return f"{'+' if (plus and x >= 0) else ''}{x:,.2f}%"


def kgm(x):
    if x is None:
        return "n.a."
    return f"{x/1e6:,.3f} million kgs"


def _word(v):
    return "increased" if (v or 0) >= 0 else "decreased"


def _word2(v):
    return "increase" if (v or 0) >= 0 else "decrease"


# ── Shared block list ────────────────────────────────────────────────────────

def build_blocks(ctx):
    B = []
    yr = ctx["year"]
    B.append(("p", "This Office respectfully submits the Comprehensive Daily Collection "
              "Report and the Import Commodity Analysis of entries received by the "
              "Assessment Division of this Port."))

    # A
    a = ctx["A"]
    B.append(("h1", f"A. Daily Collection Report as of {ctx['period_label']}"))
    B.append(("p",
        f"For the period {ctx['period_label']}, the Port of Batangas recorded a total "
        f"revenue collection of {php(a['actual'])}, "
        f"{'surpassing' if a['diff'] >= 0 else 'falling short of'} the {php(a['target'])} "
        f"DBCC Emerging Target by {php(abs(a['diff']))} or {pct(abs(a['dev_pct']), plus=False)}, "
        f"indicating {'strong' if a['diff'] >= 0 else 'weak'} revenue performance during the period."))
    B.append(("table", {"cols": ["Date Coverage", "Actual Collection", "DBCC Target", "Diff", "Dev%"],
                        "rows": [[ctx['period_label'], money(a['actual']), money(a['target']),
                                  money(a['diff']), pct(a['dev_pct'])]]}))

    # B
    b = ctx["B"]
    B.append(("h1", f"B. Cumulative Collection Report (January 1 to {ctx['end_label']})"))
    B.append(("p",
        f"From January 1 to {ctx['end_label']}, the Port of Batangas posted a total revenue "
        f"collection of {php(b['actual'])}, "
        f"{'exceeding' if b['diff'] >= 0 else 'below'} the DBCC Emerging Target of {php(b['target'])} "
        f"by {php(abs(b['diff']))} or {pct(abs(b['dev_pct']), plus=False)}, reflecting "
        f"{'positive' if b['diff'] >= 0 else 'soft'} performance over the cumulative period."))
    B.append(("table", {"cols": ["Date Coverage", "Actual Collection", "DBCC Target", "Diff", "Dev%"],
                        "rows": [[f"Jan 1 - {ctx['end_short']}, {yr}", money(b['actual']),
                                  money(b['target']), money(b['diff']), pct(b['dev_pct'])]]}))

    # C
    c = ctx["C"]
    B.append(("h1", f"C. Actual Collection vs Last Year (January 1 to {ctx['end_label']})"))
    if c["has_prior"]:
        B.append(("p",
            f"As of {ctx['end_label']}, the Port of Batangas recorded a total revenue collection of "
            f"{php(c['cur'])}, {'surpassing' if c['diff'] >= 0 else 'below'} the {php(c['prev'])} "
            f"collected during the same period in {yr-1} by {php(abs(c['diff']))}, equivalent to "
            f"{pct(abs(c['pct']), plus=False)} year-on-year {_word2(c['diff'])}. This reflects an "
            f"overall collection efficiency of {c['eff']:,.2f}%, indicating "
            f"{'strong' if c['diff'] >= 0 else 'softer'} performance relative to the previous year."))
        B.append(("table", {"cols": ["Date Coverage", f"{yr} Collection", f"{yr-1} Collection",
                                     "Diff", "YoY", "Efficiency"],
                            "rows": [[f"Jan 1 - {ctx['end_short']}", money(c['cur']), money(c['prev']),
                                      money(c['diff']), pct(c['pct']), f"{c['eff']:,.2f}%"]]}))
    else:
        B.append(("p", "Prior-year data was not loaded, so the year-on-year comparison is "
                       "unavailable. Upload a prior-year masterlist to enable this section."))

    # Daily Revenue Contributor
    dr = ctx["contrib"]
    B.append(("h2", f"Daily Revenue Contributor ({dr['label']})"))
    if dr["rows"]:
        names = dr["rows"]
        lead = (f"For {dr['label']} collection, {names[0]['name']} recorded the highest revenue "
                f"contribution at {php(names[0]['rev'])}")
        if len(names) > 1:
            lead += f", followed by {names[1]['name']} with {php(names[1]['rev'])}"
        if len(names) > 2:
            lead += f", and {names[2]['name']} with {php(names[2]['rev'])}"
        B.append(("p", lead + "."))
        rows = [[r["name"], money(r["rev"])] for r in names[:35]]
        B.append(("table", {"cols": ["Consignee", "Duties/Taxes"], "rows": rows,
                            "total": ["Grand Total", money(dr["total"])]}))

    # D
    B.append(("h1", "D. Analytics of Volume vs. Revenue"))
    d1 = ctx["D1"]
    B.append(("h2", f"D.1. {ctx['cur_month']} Volume vs. Last Month ({ctx['prev_month']})"))
    B.extend(_vol_paras(d1, ctx['prev_month'], ctx['cur_month']))
    d2 = ctx["D2"]
    B.append(("h2", f"D.2. {ctx['cur_month']} Volume vs. Last Year ({ctx['prev_year_month']})"))
    if d2["has_prior"]:
        B.extend(_vol_paras(d2, ctx['prev_year_month'], ctx['cur_month']))
    else:
        B.append(("p", "Prior-year data was not loaded, so this comparison is unavailable."))

    # E
    e = ctx["E"]
    B.append(("h1", "E. Trends of Volume and Duties of Top 30 Importers"))
    B.append(("h3", f"Compared Last Month ({ctx['prev_month']})"))
    B.extend(_trend_paras(e, "mom", "importer", ctx['period_label'], "last month"))
    B.append(("table", _trend_table(e["top30"], "Consignee", e["has_prior"])))
    B.append(("h3", f"Compared Last Year ({ctx['prev_year_month']})"))
    if e["has_prior"]:
        B.extend(_trend_paras(e, "yoy", "importer", ctx['period_label'], "last year"))
    else:
        B.append(("p", "Prior-year data was not loaded, so this comparison is unavailable."))

    # F
    f = ctx["F"]
    B.append(("h1", "F. Trends of Volume and Duties of Top 30 Commodities"))
    B.append(("h3", f"Compared Last Month ({ctx['prev_month']})"))
    B.extend(_comm_paras(f, "mom", ctx['period_label'], "last month"))
    B.append(("table", _trend_table(f["top30"], "Commodity", f["has_prior"])))
    B.append(("h3", f"Compared Last Year ({ctx['prev_year_month']})"))
    if f["has_prior"]:
        B.extend(_comm_paras(f, "yoy", ctx['period_label'], "last year"))
    else:
        B.append(("p", "Prior-year data was not loaded, so this comparison is unavailable."))

    return B


def _vol_paras(d, prev_label, cur_label):
    if not d.get("has", True) or d["prev_vol"] == 0:
        return [("p", f"Total import volume for {cur_label} was {kgm(d['cur_vol'])} "
                 f"(no comparable {prev_label} volume on record).")]
    out = [("p",
        f"Total import volume {_word(d['vol_pct'])} by {pct(abs(d['vol_pct']), plus=False)} from "
        f"{kgm(d['prev_vol'])} in {prev_label} to {kgm(d['cur_vol'])} in {cur_label}, primarily "
        f"driven by {pct(d['nonoil_vol_pct'])} and {pct(d['oil_vol_pct'])} change in non-oil and oil "
        f"imports, respectively. The {_word2(d['nonoil_vol_pct'])} in non-oil imports resulted in a "
        f"{php(abs(d['nonoil_rev_d']))} or {pct(abs(d['nonoil_rev_pct']), plus=False)} "
        f"{_word2(d['nonoil_rev_d'])} in revenue, while the {_word2(d['oil_vol_pct'])} in oil import "
        f"volume translated into {php(abs(d['oil_rev_d']))} or {pct(abs(d['oil_rev_pct']), plus=False)} "
        f"{_word2(d['oil_rev_d'])} in revenue.")]
    out.append(("p",
        f"Overall, the change in revenue from both oil and non-oil commodities resulted in a net "
        f"revenue {_word2(d['net_rev_d'])} of {php(abs(d['net_rev_d']))}, equivalent to "
        f"{pct(abs(d['net_rev_pct']), plus=False)}. In other words, the movement in both oil and "
        f"non-oil import volumes had a substantial impact on revenue performance for the period."))
    return out


def _trend_paras(e, mode, noun, period_label, vs_label):
    new = e[f"{mode}_new"]
    grow = e[f"{mode}_grow"]
    out = [("p",
        f"For the period {period_label}, compared to the same period {vs_label}, overall revenue "
        f"collections show a significant increase in both revenue and volume among several major "
        f"importers. Despite some decreases, the overall growth is largely driven by the sharp "
        f"increases recorded by top importers.")]
    if new:
        s = ("Notably, several importers lodged their entries this period but did not do so for the "
             "same period " + vs_label + ". Among them, ")
        parts = [f"{n['name']} reported {php(n['rev'])} (+100%) in revenue and {kgm(n['vol'])} "
                 f"(+100%) in volume" for n in new[:3]]
        s += "; ".join(parts) + "."
        out.append(("p", s))
    if grow:
        s = "Meanwhile, certain importers also posted increases, including "
        parts = [f"{g['name']} with {php(g['rev_d'])} or {pct(g['rev_pct'])} in revenue and "
                 f"{kgm(g['vol_d'])} or {pct(g['vol_pct'])} in volume" for g in grow[:3]]
        s += "; ".join(parts) + "."
        out.append(("p", s))
    out.append(("p", "These increases highlight the impact of POB's top importers on the overall "
                     "revenue collection for the period."))
    return out


def _comm_paras(f, mode, period_label, vs_label):
    top = f[f"{mode}_top"]
    out = [("p", f"For the period {period_label}, several top commodities registered significant "
            f"growth in both revenue and volume compared to the same period {vs_label}.")]
    if top:
        lead = top[0]
        s = (f"In particular, {lead['name']} recorded the highest revenue increase of "
             f"{php(lead['rev_d'])} ({pct(lead['rev_pct'])}) and volume change of "
             f"{kgm(lead['vol_d'])} ({pct(lead['vol_pct'])}).")
        if len(top) > 1:
            s += (f" {top[1]['name']} followed with a revenue change of {php(top[1]['rev_d'])} "
                  f"({pct(top[1]['rev_pct'])}) and {kgm(top[1]['vol_d'])} in volume.")
        if len(top) > 2:
            s += (f" Similarly, {top[2]['name']} reflected a revenue change of {php(top[2]['rev_d'])} "
                  f"({pct(top[2]['rev_pct'])}).")
        out.append(("p", s))
    drivers = f.get("drivers", [])
    if drivers:
        out.append(("p",
            f"The change in these commodities can be observed with the revenue collection from top "
            f"importers such as {', '.join(drivers[:3])}, all of which registered substantial "
            f"movement in both revenue and volume. This significantly affects the overall collection "
            f"of the Port during the said period."))
    return out


def _trend_table(rows, noun, has_prior=True):
    cols = ["#", noun, "Revenue", "Volume (kgs)", "ΔRev MoM", "ΔVol MoM", "ΔRev YoY", "ΔVol YoY"]
    out = []
    for i, r in enumerate(rows, 1):
        yr = pct(r["rev_yoy"]) if has_prior else "—"
        yv = pct(r["vol_yoy"]) if has_prior else "—"
        out.append([str(i), r["name"], money(r["revenue"]), money(r["volume"]),
                    pct(r["rev_mom"]), pct(r["vol_mom"]), yr, yv])
    return {"cols": cols, "rows": out, "small": True}


# ── HTML rendering (Memo tab) ────────────────────────────────────────────────

def _esc(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def narrative_html(ctx):
    o = OFFICIALS
    h = ['<div class="memo-doc">']
    h.append('<div class="memo-head"><div class="memo-title">MEMORANDUM</div>'
             f'<div class="memo-ref">{_esc(o["ref"])}</div></div>')
    h.append('<table class="memo-fields"><tbody>')
    for lbl, nm, ti in (("FOR", o["for_name"], o["for_title"]),
                        ("THRU", o["thru_name"], o["thru_title"]),
                        ("FROM", o["from_name"], o["from_title"])):
        h.append(f'<tr><td class="mf-l">{lbl}</td><td><strong>{_esc(nm)}</strong><br>'
                 f'<span class="mf-t">{_esc(ti)}</span></td></tr>')
    h.append(f'<tr><td class="mf-l">SUBJECT</td><td>{_esc(o["subject"])}<br>'
             f'<span class="mf-t">Summary Report as of {_esc(ctx["end_label"])}</span></td></tr>')
    h.append(f'<tr><td class="mf-l">DATE</td><td>{_esc(ctx["date_str"])}</td></tr>')
    h.append('</tbody></table><hr>')
    for blk in build_blocks(ctx):
        k = blk[0]
        if k == "h1":
            h.append(f'<h2>{_esc(blk[1])}</h2>')
        elif k == "h2":
            h.append(f'<h3>{_esc(blk[1])}</h3>')
        elif k == "h3":
            h.append(f'<h4>{_esc(blk[1])}</h4>')
        elif k == "p":
            h.append(f'<p>{_esc(blk[1])}</p>')
        elif k == "table":
            h.append(_html_table(blk[1]))
    h.append('<p style="margin-top:1.5rem">Respectfully submitted.</p>')
    h.append('</div>')
    return "\n".join(h)


def _html_table(t):
    cls = "memo-tbl small" if t.get("small") else "memo-tbl"
    h = [f'<div class="memo-tbl-wrap"><table class="{cls}"><thead><tr>']
    for c in t["cols"]:
        h.append(f'<th>{_esc(c)}</th>')
    h.append('</tr></thead><tbody>')
    for row in t["rows"]:
        h.append('<tr>' + ''.join(f'<td>{_esc(x)}</td>' for x in row) + '</tr>')
    if t.get("total"):
        h.append('<tr class="tot">' + ''.join(f'<td>{_esc(x)}</td>' for x in t["total"]) + '</tr>')
    h.append('</tbody></table></div>')
    return "".join(h)


# ── DOCX rendering ───────────────────────────────────────────────────────────

def _shade(cell, hexc):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexc)
    cell._tc.get_or_add_tcPr().append(sh)


def _set_cell(cell, text, *, bold=False, white=False, align="left", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_table(doc, t):
    cols = t["cols"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    small = t.get("small")
    sz = 7.5 if small else 9
    # header
    for j, c in enumerate(cols):
        cell = tbl.rows[0].cells[j]
        _shade(cell, NAVY)
        _set_cell(cell, c, bold=True, white=True,
                  align="left" if j <= 1 else "right" if j else "center", size=sz)
    numeric_from = 2 if not (cols and cols[0] == "#") else 2
    for row in t["rows"]:
        cells = tbl.add_row().cells
        for j, val in enumerate(row):
            al = "left" if (j == 0 and cols[0] != "#") or (j == 1 and cols[0] == "#") else \
                 ("center" if (cols[0] == "#" and j == 0) else "right")
            _set_cell(cells[j], val, align=al, size=sz)
    if t.get("total"):
        cells = tbl.add_row().cells
        for j, val in enumerate(t["total"]):
            _shade(cells[j], LIGHT)
            _set_cell(cells[j], val, bold=True, align="left" if j == 0 else "right", size=sz)
    doc.add_paragraph()


def build_docx(ctx):
    o = OFFICIALS
    # Use the bundled template (carries the original letterhead header + footer)
    # when present; fall back to a blank document otherwise.
    tmpl = os.path.join(os.path.dirname(os.path.abspath(__file__)), "memo_template.docx")
    doc = Document(tmpl) if os.path.exists(tmpl) else Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    def title(txt, size=14, align="left"):
        p = doc.add_paragraph()
        p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(size)
        return p

    title("MEMORANDUM", 14)
    doc.add_paragraph(o["ref"])
    doc.add_paragraph()

    # FOR / THRU / FROM / SUBJECT / DATE
    def field(label, name, sub=None):
        p = doc.add_paragraph()
        r = p.add_run(f"{label}\t:\t"); r.bold = True
        r2 = p.add_run(name); r2.bold = True
        if sub:
            ps = doc.add_paragraph("\t\t" + sub)
            ps.paragraph_format.space_after = Pt(2)
    field("FOR", o["for_name"], o["for_title"])
    field("THRU", o["thru_name"], o["thru_title"])
    field("FROM", o["from_name"], o["from_title"])
    p = doc.add_paragraph()
    r = p.add_run("SUBJECT\t:\t"); r.bold = True
    p.add_run(o["subject"])
    doc.add_paragraph("\t\tSummary Report as of " + ctx["end_label"])
    p = doc.add_paragraph()
    r = p.add_run("DATE\t:\t"); r.bold = True
    p.add_run(ctx["date_str"])
    doc.add_paragraph("_" * 70)

    for blk in build_blocks(ctx):
        k = blk[0]
        if k == "h1":
            doc.add_paragraph()
            title(blk[1], 12)
        elif k == "h2":
            title(blk[1], 11)
        elif k == "h3":
            p = doc.add_paragraph(); r = p.add_run(blk[1]); r.bold = True; r.italic = True
        elif k == "p":
            p = doc.add_paragraph(blk[1])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif k == "table":
            _add_table(doc, blk[1])

    doc.add_paragraph()
    doc.add_paragraph("Respectfully submitted.")

    # Keep the section properties (which hold the header/footer references) as the
    # final body element so Word renders the letterhead header + footer correctly.
    body = doc.element.body
    sp = body.find(qn('w:sectPr'))
    if sp is not None:
        body.remove(sp)
        body.append(sp)
        # Fix zero top margin — the template has w:top="0" which causes body text
        # to overlap the letterhead header image (1.76 inches tall, ~2530 DXA).
        # Enforce at least 1.8 inches (2592 DXA) to clear the header.
        pgmar = sp.find(qn('w:pgMar'))
        if pgmar is not None:
            top = int(pgmar.get(qn('w:top'), '0'))
            if top < 2592:
                pgmar.set(qn('w:top'), '2592')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
